#!/usr/bin/evn python
#-*- coding:utf-8 -*-
'''
Created on 2018年7月12日

@author: jinfeng
'''
from docx import Document
from docx.shared import Inches

document = Document('demo.docx')
for paragraph in document.paragraphs:
    print(paragraph.text)
    if 'Add new paragraph' in paragraph.text:
        paragraph.delete

document.add_paragraph(
    'Add new paragraph', style='ListNumber'
)    

document.save('demo.docx')

